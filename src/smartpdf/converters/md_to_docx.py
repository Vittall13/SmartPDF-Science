"""Markdown to DOCX converter with formula support."""
import re
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)


class MarkdownToDOCX:
    """Convert Markdown to Microsoft Word DOCX format."""
    
    def __init__(self):
        self.doc = None
    
    def convert(self, markdown_text: str, output_path: str | Path) -> None:
        """Convert Markdown to DOCX.
        
        Args:
            markdown_text: Markdown text to convert
            output_path: Path to save DOCX file
        """
        logger.info(f"Converting to DOCX: {output_path}")
        
        self.doc = Document()
        lines = markdown_text.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines and comments
            if not line or line.startswith('<!--'):
                i += 1
                continue
            
            # Headers
            if line.startswith('#'):
                level = len(line.split()[0])
                text = line.lstrip('#').strip()
                heading = self.doc.add_heading(text, level=min(level, 3))
            
            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                self.doc.add_paragraph(text, style='List Bullet')
            
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                self.doc.add_paragraph(text, style='List Number')
            
            # Tables
            elif line.startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_table(table_lines)
                continue
            
            # Formulas (LaTeX blocks)
            elif line.startswith('$$'):
                formula_lines = [line]
                i += 1
                while i < len(lines) and not lines[i].strip().endswith('$$'):
                    formula_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    formula_lines.append(lines[i])
                formula = '\n'.join(formula_lines)
                self._add_formula(formula)
            
            # Regular paragraphs
            else:
                self.doc.add_paragraph(line)
            
            i += 1
        
        self.doc.save(str(output_path))
        logger.info(f"DOCX saved: {output_path}")
    
    def _add_table(self, table_lines: list) -> None:
        """Add table to document."""
        # Parse table
        rows = []
        for line in table_lines:
            if '---' in line:  # Skip separator line
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)
        
        if not rows:
            return
        
        # Create table
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'
        
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                row.cells[j].text = cell_text
    
    def _add_formula(self, formula: str) -> None:
        """Add formula as text (Word equation would require OMML conversion)."""
        # Remove $$ markers
        formula = formula.replace('$$', '').strip()
        
        # Add as monospace paragraph
        para = self.doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(formula)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)